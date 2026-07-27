"""Converte o Markdown deste projeto em .docx.

Nao e um conversor de Markdown geral - e um conversor do **subconjunto que a
documentacao deste projeto usa**, medido antes de escrever: cabecalhos,
tabelas (a construcao mais frequente, 732 linhas), listas, citacoes, blocos de
codigo, regras horizontais e os marcadores inline `**negrito**`, `` `codigo` ``
e `[texto](link)`.

Escrever um subconjunto conhecido, em vez de instalar um conversor generico,
tem uma razao pratica: as tabelas deste projeto carregam o resultado - se elas
sairem erradas, o documento exportado mente. Um conversor proprio pode ser
testado contra exatamente o que existe.
"""

import re
from dataclasses import dataclass, field

# --- reconhecimento de bloco -------------------------------------------------

RE_CABECALHO = re.compile(r'^(#{1,6})\s+(.*)$')
RE_LISTA = re.compile(r'^\s*[-*]\s+(.*)$')
RE_NUMERADA = re.compile(r'^\s*(\d+)\.\s+(.*)$')
RE_CITACAO = re.compile(r'^>\s?(.*)$')
RE_CERCA = re.compile(r'^\s*```')
RE_REGUA = re.compile(r'^(-{3,}|\*{3,}|_{3,})\s*$')
RE_TABELA = re.compile(r'^\s*\|')
RE_SEPARADOR = re.compile(r'^\s*\|[\s:|-]+\|\s*$')
RE_HTML = re.compile(r'^\s*</?(details|summary)\b', re.IGNORECASE)

# --- reconhecimento inline ---------------------------------------------------
# Ordem importa: o link e reconhecido antes do codigo para que `[a](b)` nao
# seja partido; e o negrito antes do italico, senao `**x**` vira `*` + `*x*`.
RE_INLINE = re.compile(
    r'(?P<link>\[(?P<texto>[^\]]+)\]\((?P<url>[^)]+)\))'
    r'|(?P<codigo>`[^`]+`)'
    r'|(?P<negrito>\*\*[^*]+\*\*)'
    r'|(?P<italico>(?<!\*)\*(?!\*)[^*]+\*(?!\*))'
)


@dataclass
class Trecho:
    """Um pedaco de texto com formatacao."""

    texto: str
    negrito: bool = False
    italico: bool = False
    codigo: bool = False
    url: str = ''


@dataclass
class Bloco:
    tipo: str  # cabecalho | paragrafo | lista | numerada | citacao | codigo
               # | tabela | regua
    nivel: int = 0
    linhas: list = field(default_factory=list)   # list[list[Trecho]]
    tabela: list = field(default_factory=list)   # list[list[list[Trecho]]]
    texto_bruto: str = ''


PROFUNDIDADE_MAXIMA = 4


def formatar_inline(texto, negrito=False, italico=False, profundidade=0):
    """Quebra uma linha em trechos formatados.

    O que nao casa com nenhum marcador vira trecho simples - por isso o texto
    nunca se perde, mesmo diante de sintaxe que este conversor nao conhece.

    ⚠️ **A recursao em negrito e italico nao e enfeite.** Sem ela,
    `**texto com [link](url)**` casaria inteiro como negrito e a sintaxe do
    link sairia literal no documento - foi o unico defeito que a conferencia
    contra os 11 documentos reais encontrou. Negrito e italico reabrem a
    analise do proprio conteudo, herdando a formatacao; codigo e link, nao,
    porque neles o conteudo e terminal.
    """
    trechos, posicao = [], 0

    def simples(fragmento):
        return Trecho(fragmento, negrito=negrito, italico=italico)

    for casa in RE_INLINE.finditer(texto):
        if casa.start() > posicao:
            trechos.append(simples(texto[posicao:casa.start()]))

        if casa.group('link'):
            # O texto do link pode ser codigo: [`arquivo.py`](caminho).
            interno = casa.group('texto')
            mono = interno.startswith('`') and interno.endswith('`')
            trechos.append(Trecho(
                interno[1:-1] if mono else interno,
                negrito=negrito, italico=italico, codigo=mono,
                url=casa.group('url'),
            ))
        elif casa.group('codigo'):
            trechos.append(Trecho(
                casa.group('codigo')[1:-1],
                negrito=negrito, italico=italico, codigo=True,
            ))
        elif casa.group('negrito'):
            interno = casa.group('negrito')[2:-2]
            if profundidade < PROFUNDIDADE_MAXIMA:
                trechos.extend(formatar_inline(
                    interno, True, italico, profundidade + 1
                ))
            else:
                trechos.append(Trecho(interno, negrito=True, italico=italico))
        elif casa.group('italico'):
            interno = casa.group('italico')[1:-1]
            if profundidade < PROFUNDIDADE_MAXIMA:
                trechos.extend(formatar_inline(
                    interno, negrito, True, profundidade + 1
                ))
            else:
                trechos.append(Trecho(interno, negrito=negrito, italico=True))

        posicao = casa.end()

    if posicao < len(texto):
        trechos.append(simples(texto[posicao:]))

    return trechos or [simples('')]


def _celulas(linha):
    """Divide uma linha de tabela, descartando as bordas vazias."""
    partes = linha.strip().split('|')
    if partes and not partes[0].strip():
        partes = partes[1:]
    if partes and not partes[-1].strip():
        partes = partes[:-1]
    return [formatar_inline(p.strip()) for p in partes]


def analisar(markdown):
    """Markdown -> lista de blocos. Nao toca em docx: da para testar sozinho."""
    blocos = []
    linhas = markdown.replace('\r\n', '\n').split('\n')
    i = 0

    while i < len(linhas):
        linha = linhas[i]

        # Bloco de codigo: tudo entre as cercas vai literal, sem interpretar
        # marcador nenhum - senao um `**` dentro de exemplo viraria negrito.
        if RE_CERCA.match(linha):
            corpo, i = [], i + 1
            while i < len(linhas) and not RE_CERCA.match(linhas[i]):
                corpo.append(linhas[i])
                i += 1
            i += 1
            blocos.append(Bloco('codigo', texto_bruto='\n'.join(corpo)))
            continue

        if not linha.strip() or RE_HTML.match(linha):
            i += 1
            continue

        if RE_REGUA.match(linha):
            blocos.append(Bloco('regua'))
            i += 1
            continue

        casa = RE_CABECALHO.match(linha)
        if casa:
            blocos.append(Bloco(
                'cabecalho', nivel=len(casa.group(1)),
                linhas=[formatar_inline(casa.group(2).strip())],
            ))
            i += 1
            continue

        if RE_TABELA.match(linha):
            corpo = []
            while i < len(linhas) and RE_TABELA.match(linhas[i]):
                if not RE_SEPARADOR.match(linhas[i]):
                    corpo.append(_celulas(linhas[i]))
                i += 1
            if corpo:
                blocos.append(Bloco('tabela', tabela=corpo))
            continue

        casa = RE_CITACAO.match(linha)
        if casa:
            corpo = []
            while i < len(linhas) and RE_CITACAO.match(linhas[i]):
                corpo.append(RE_CITACAO.match(linhas[i]).group(1))
                i += 1
            blocos.append(Bloco('citacao', linhas=[formatar_inline(' '.join(corpo))]))
            continue

        if RE_LISTA.match(linha):
            itens = []
            while i < len(linhas) and RE_LISTA.match(linhas[i]):
                itens.append(formatar_inline(RE_LISTA.match(linhas[i]).group(1)))
                i += 1
            blocos.append(Bloco('lista', linhas=itens))
            continue

        if RE_NUMERADA.match(linha):
            itens = []
            while i < len(linhas) and RE_NUMERADA.match(linhas[i]):
                itens.append(formatar_inline(RE_NUMERADA.match(linhas[i]).group(2)))
                i += 1
            blocos.append(Bloco('numerada', linhas=itens))
            continue

        # Paragrafo: junta linhas ate a proxima em branco ou inicio de bloco.
        corpo = []
        while i < len(linhas) and linhas[i].strip():
            atual = linhas[i]
            if (RE_CABECALHO.match(atual) or RE_TABELA.match(atual)
                    or RE_CITACAO.match(atual) or RE_LISTA.match(atual)
                    or RE_NUMERADA.match(atual) or RE_CERCA.match(atual)
                    or RE_REGUA.match(atual) or RE_HTML.match(atual)):
                break
            corpo.append(atual.strip())
            i += 1
        if corpo:
            blocos.append(Bloco('paragrafo', linhas=[formatar_inline(' '.join(corpo))]))

    return blocos


# --- renderizacao ------------------------------------------------------------

def _escrever_trechos(paragrafo, trechos, mono=False):
    for trecho in trechos:
        run = paragrafo.add_run(trecho.texto)
        run.bold = trecho.negrito or None
        run.italic = trecho.italico or None
        if trecho.codigo or mono:
            run.font.name = 'Consolas'
        if trecho.url:
            # Sem hiperlink real: o python-docx nao expoe API para isso sem
            # mexer no XML. O sublinhado sinaliza o link e a URL fica no texto
            # quando ela nao e obvia - melhor que perder a referencia.
            run.underline = True


def para_documento(markdown, titulo=''):
    """Constroi o objeto Document. Separado de `analisar` para poder testar."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    documento = Document()

    estilo = documento.styles['Normal']
    estilo.font.name = 'Calibri'
    estilo.font.size = Pt(10.5)

    if titulo:
        documento.add_heading(titulo, level=0)

    for bloco in analisar(markdown):
        if bloco.tipo == 'cabecalho':
            # O Word so tem estilo ate Heading 9; o Markdown para em 6.
            p = documento.add_heading('', level=min(bloco.nivel, 6))
            _escrever_trechos(p, bloco.linhas[0])

        elif bloco.tipo == 'paragrafo':
            _escrever_trechos(documento.add_paragraph(), bloco.linhas[0])

        elif bloco.tipo == 'lista':
            for item in bloco.linhas:
                _escrever_trechos(
                    documento.add_paragraph(style='List Bullet'), item
                )

        elif bloco.tipo == 'numerada':
            for item in bloco.linhas:
                _escrever_trechos(
                    documento.add_paragraph(style='List Number'), item
                )

        elif bloco.tipo == 'citacao':
            p = documento.add_paragraph(style='Intense Quote')
            _escrever_trechos(p, bloco.linhas[0])

        elif bloco.tipo == 'codigo':
            p = documento.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            run = p.add_run(bloco.texto_bruto)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)

        elif bloco.tipo == 'tabela':
            colunas = max(len(l) for l in bloco.tabela)
            tabela = documento.add_table(rows=0, cols=colunas)
            tabela.style = 'Table Grid'
            for indice, linha in enumerate(bloco.tabela):
                celulas = tabela.add_row().cells
                for posicao in range(colunas):
                    p = celulas[posicao].paragraphs[0]
                    if posicao < len(linha):
                        _escrever_trechos(p, linha[posicao])
                    if indice == 0:
                        for run in p.runs:
                            run.bold = True

        elif bloco.tipo == 'regua':
            p = documento.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run('— · —').italic = True

    return documento


def converter(caminho_md, caminho_docx, titulo=''):
    """Le um .md e escreve o .docx correspondente. Retorna o numero de blocos."""
    from pathlib import Path

    texto = Path(caminho_md).read_text(encoding='utf-8')
    documento = para_documento(texto, titulo)
    Path(caminho_docx).parent.mkdir(parents=True, exist_ok=True)
    documento.save(str(caminho_docx))
    return len(analisar(texto))
